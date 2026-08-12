"""Sales / stock statement extraction into a unified JSON schema.

Separate from invoice OCR — do not use for tax invoices.
"""

from __future__ import annotations

import base64
import io
import json
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {
    ".txt",
    ".htm",
    ".html",
    ".doc",
    ".docx",
    ".xls",
    ".xlsx",
    ".pdf",
    ".jpg",
    ".jpeg",
    ".png",
    ".bmp",
    ".tif",
    ".tiff",
    ".webp",
}

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff", ".webp"}
TEXT_EXTENSIONS = {".txt"}
WORD_EXTENSIONS = {".doc", ".docx"}


_MONTH_MAP = {
    "jan": 1,
    "january": 1,
    "feb": 2,
    "february": 2,
    "mar": 3,
    "march": 3,
    "apr": 4,
    "april": 4,
    "may": 5,
    "jun": 6,
    "june": 6,
    "jul": 7,
    "july": 7,
    "aug": 8,
    "august": 8,
    "sep": 9,
    "sept": 9,
    "september": 9,
    "oct": 10,
    "october": 10,
    "nov": 11,
    "november": 11,
    "dec": 12,
    "december": 12,
}


def empty_result(source_file: str = "", source_format: str = "") -> Dict[str, Any]:
    return {
        "source_file": source_file,
        "source_format": source_format,
        "stockist_name": None,
        "stockist_address": None,
        "company_name": None,
        "period_from": None,
        "period_to": None,
        "report_title": None,
        "line_items": [],
        "totals": {
            "sales_value": None,
            "closing_value": None,
            "extra": {},
        },
    }


def empty_line_item() -> Dict[str, Any]:
    return {
        "product_code": None,
        "product_name": None,
        "packing": None,
        "opening_qty": 0.0,
        "receipts_qty": 0.0,
        "sales_qty": 0.0,
        "sales_value": 0.0,
        "closing_qty": 0.0,
        "closing_value": 0.0,
        "extra": {},
    }


def _to_float(value: Any, default: float = 0.0) -> float:
    if value is None:
        return default
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).strip().replace(",", "").replace("L", "").replace("l", "")
    if not text or text in {"-", "—", "NA", "N/A"}:
        return default
    try:
        return float(text)
    except ValueError:
        m = re.search(r"-?\d+(?:\.\d+)?", text)
        return float(m.group(0)) if m else default


def _to_nullable_float(value: Any) -> Optional[float]:
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None
    return _to_float(value, 0.0)


def _normalize_date(raw: Optional[str]) -> Optional[str]:
    if not raw:
        return None
    text = str(raw).strip()
    if not text:
        return None

    # 01/06/2026, 01-06-2026, 01.06.2026
    m = re.search(r"(\d{1,2})[./-](\d{1,2})[./-](\d{2,4})", text)
    if m:
        d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if y < 100:
            y += 2000
        try:
            return datetime(y, mo, d).strftime("%Y-%m-%d")
        except ValueError:
            pass

    # 01-Jun-26 / 01-Jun-2026
    m = re.search(r"(\d{1,2})[- ]([A-Za-z]{3,9})[- ](\d{2,4})", text)
    if m:
        d = int(m.group(1))
        mo = _MONTH_MAP.get(m.group(2).lower()[:3]) or _MONTH_MAP.get(m.group(2).lower())
        y = int(m.group(3))
        if y < 100:
            y += 2000
        if mo:
            try:
                return datetime(y, mo, d).strftime("%Y-%m-%d")
            except ValueError:
                pass

    # Month Of Jul 2026 → first/last day handled by caller period helpers
    return None


def _month_period(month_name: str, year: int) -> Tuple[Optional[str], Optional[str]]:
    mo = _MONTH_MAP.get(month_name.lower()[:3]) or _MONTH_MAP.get(month_name.lower())
    if not mo:
        return None, None
    start = datetime(year, mo, 1)
    if mo == 12:
        end = datetime(year, 12, 31)
    else:
        end = datetime(year, mo + 1, 1)
        from datetime import timedelta

        end = end - timedelta(days=1)
    return start.strftime("%Y-%m-%d"), end.strftime("%Y-%m-%d")


def _clean_name(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip())


def _ocr_image_footer_text(image_bytes: bytes, top_ratio: float = 0.78) -> str:
    """OCR only the bottom band of a page image to capture TOTAL rows."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return ""

    import os

    cmd = os.getenv("TESSERACT_CMD", "").strip()
    if cmd:
        pytesseract.pytesseract.tesseract_cmd = cmd
    elif os.name == "nt":
        win_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        if os.path.exists(win_cmd):
            pytesseract.pytesseract.tesseract_cmd = win_cmd

    try:
        img = Image.open(io.BytesIO(image_bytes))
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        w, h = img.size
        crop = img.crop((0, int(h * top_ratio), w, h))
        return pytesseract.image_to_string(crop, config="--psm 6") or ""
    except Exception as exc:
        logger.warning("Footer OCR failed: %s", exc)
        return ""


def _merge_footer_total_into_text(page_text: str, image_bytes: Optional[bytes]) -> str:
    """Ensure TOTAL footer line is present in page text (full-page OCR often truncates it)."""
    text = page_text or ""
    if re.search(r"^\s*TOTAL\b.*\d", text, re.I | re.M):
        # Already has a TOTAL with numbers
        return text
    if not image_bytes:
        return text
    footer = _ocr_image_footer_text(image_bytes)
    if not footer.strip():
        return text
    total_lines = [
        ln.strip()
        for ln in footer.splitlines()
        if re.match(r"^\s*TOTAL\b", ln.strip(), re.I)
    ]
    if not total_lines:
        # keep whole footer band — may still help parsers
        return text + "\n" + footer
    return text + "\n" + "\n".join(total_lines)


def _parse_ps_pharma_statement(text: str, filename: str) -> Optional[Dict[str, Any]]:
    """Parse P.S.PHARMACEUTICALS OPENING/RECEIPT/ISSUE/CLOSING stock & sales analysis."""
    if not text:
        return None
    # Only P.S. format (OPENING/RECEIPT/ISSUE/CLOSING) — do not steal Mahajan pages
    is_ps = bool(
        re.search(r"P\.?\s*S\.?\s*PHARMACEUTICAL", text, re.I)
        and (
            re.search(r"OPENING\s+RECEIPT\s+ISSUE\s+CLOSING", text, re.I)
            or re.search(r"STOCK\s*&\s*SALES\s+ANALYSIS", text, re.I)
            or re.search(r"^\s*TOTAL\b", text, re.I | re.M)
        )
    ) or bool(re.search(r"OPENING\s+RECEIPT\s+ISSUE\s+CLOSING", text, re.I))
    if not is_ps:
        return None

    result = empty_result(filename, "pdf")
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if lines:
        result["stockist_name"] = _clean_name(lines[0])
    for ln in lines[:12]:
        if re.search(r"STOCK\s*&\s*SALES|Sales\s*&\s*Stock", ln, re.I):
            result["report_title"] = _clean_name(ln)
            m = re.search(
                r"(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})\s*[-–]+\s*(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})",
                ln,
            )
            if m:
                result["period_from"] = _normalize_date(m.group(1))
                result["period_to"] = _normalize_date(m.group(2))
        if re.search(r"AUROBINDO|VERITAZ|HEALTHCARE", ln, re.I) and not re.search(
            r"PHARMACEUTICALS|ANALYSIS", ln, re.I
        ):
            # company often appears as division line or inside title parentheses
            m_co = re.search(r"\(([^)]+)\)", ln)
            if m_co:
                result["company_name"] = _clean_name(m_co.group(1))
            elif not result.get("company_name"):
                result["company_name"] = _clean_name(ln)
        if re.search(r"RAGHUNATH|JAMMU|Phone|GSTIN", ln, re.I) and not result.get(
            "stockist_address"
        ):
            if not re.search(r"Phone|GSTIN|E-Mail|VAT", ln, re.I):
                result["stockist_address"] = _clean_name(ln)

    # Row: <product [packing]> opening receipt issue closing
    row_re = re.compile(
        r"^(.+?)\s+(-?\d+(?:\.\d+)?)\s+(-?\d+(?:\.\d+)?)\s+"
        r"(-?\d+(?:\.\d+)?)\s+(-?\d+(?:\.\d+)?)\s*$"
    )
    pack_token = re.compile(
        r"^(?:\d+\*\d+|\d+(?:\.\d+)?(?:ML|MG|TAB|CAP|SYP|S)|"
        r"\d+ML|\d+MG|0\.\d+ML)$",
        re.I,
    )

    items: List[Dict[str, Any]] = []
    for ln in lines:
        if re.match(r"^\s*TOTAL\b", ln, re.I):
            continue
        if re.search(
            r"^(ITEM|OPENING|STOCK|AUROBINDO|Phone|GSTIN|Page|P\.?S\.?)",
            ln,
            re.I,
        ):
            continue
        m = row_re.match(ln)
        if not m:
            continue
        left = _clean_name(m.group(1))
        # Skip manufacturer-only divider rows
        if re.fullmatch(r"AUROBINDO.*|VERITAZ.*", left, re.I):
            continue
        tokens = left.split()
        packing = None
        product_name = left
        if len(tokens) >= 2 and pack_token.match(tokens[-1]):
            packing = tokens[-1]
            product_name = " ".join(tokens[:-1]).strip()

        item = empty_line_item()
        item["product_name"] = product_name
        item["packing"] = packing
        item["opening_qty"] = _to_float(m.group(2))
        item["receipts_qty"] = _to_float(m.group(3))
        item["sales_qty"] = _to_float(m.group(4))  # ISSUE
        item["closing_qty"] = _to_float(m.group(5))
        item["sales_value"] = 0.0
        item["closing_value"] = 0.0
        items.append(item)

    if not items and not re.search(r"^\s*TOTAL\b", text, re.I | re.M):
        return None

    result["line_items"] = items
    result["totals"]["extra"]["extraction_method"] = "ps_pharma_parser"
    result = _apply_total_row_to_result(result, text)

    # If footer TOTAL still missing, fall back to line sums for qty totals
    totals = result["totals"]
    if totals.get("opening_qty") is None and items:
        totals["opening_qty"] = sum(_to_float(i["opening_qty"]) for i in items)
        totals["receipts_qty"] = sum(_to_float(i["receipts_qty"]) for i in items)
        totals["sales_qty"] = sum(_to_float(i["sales_qty"]) for i in items)
        totals["closing_qty"] = sum(_to_float(i["closing_qty"]) for i in items)
        totals["extra"]["total_row_source"] = "line_sum_fallback"
        totals["extra"]["total_row_labels"] = [
            "OPENING",
            "RECEIPT",
            "ISSUE",
            "CLOSING",
        ]
    return result


def _is_opbal_issue_closing_format(text: str) -> bool:
    """Mahajan-style: PRODUCT/PACKING OpBal Receipt Total Issue Closing."""
    if not text:
        return False
    header = "\n".join(ln for ln in text.splitlines()[:50] if ln.strip())
    has_op = bool(re.search(r"Op\.?\s*Bal", header, re.I))
    has_receipt = bool(re.search(r"\bReceipt\b", header, re.I))
    has_issue = bool(re.search(r"\bIssue\b", header, re.I))
    has_closing = bool(re.search(r"\bClosing\b", header, re.I))
    return has_op and has_receipt and has_issue and has_closing


def _opbal_layout(text: str) -> str:
    """Return qty column layout for OpBal statements.

    - issue_expiry_closing: OpBal Receipt Total Issue Expiry Closing Near
      (MAHAJAN & MAHAJAN)
    - issue_closing_dump: OpBal Receipt Total Issue Closing Dump Near
      (MAHAJAN ASSOCIATES)
    """
    header = "\n".join(ln for ln in (text or "").splitlines()[:50] if ln.strip())
    # Expiry/Breakage column sits between Issue and Closing
    if re.search(
        r"Issue\s+Expiry\s+Closing|Issue\s+\S*\s*Expiry\s+\S*\s*Closing|"
        r"Expiry\s+Closing\s+Near|Breakage\s+Balance",
        header,
        re.I,
    ):
        return "issue_expiry_closing"
    if re.search(r"\bExpiry\b", header, re.I) and re.search(
        r"Issue.{0,20}Expiry.{0,20}Closing", header, re.I | re.S
    ):
        return "issue_expiry_closing"
    return "issue_closing_dump"


def _ocr_qty_token(tok: str) -> Optional[float]:
    """Parse a trailing qty token; tolerate OCR letter glue (S54, a2) and o/° as 0."""
    if tok is None:
        return None
    t = str(tok).strip().replace(",", "")
    if not t:
        return None
    if t in {"o", "O", "°", "〇", "()", "—", "-", "."}:
        return 0.0
    if re.fullmatch(r"-?\d+(?:\.\d+)?", t):
        return _to_float(t)
    # Single letter glued onto digits (common Tesseract noise on this format)
    m = re.fullmatch(r"[A-Za-z](-?\d+(?:\.\d+)?)", t)
    if m:
        return _to_float(m.group(1))
    # Trailing junk letter: 66o, 30_
    m = re.fullmatch(r"(-?\d+(?:\.\d+)?)[A-Za-z._]*", t)
    if m and not re.search(r"(?:TAB|ML|MG|CAP|SYP|INJ)\b", t, re.I):
        return _to_float(m.group(1))
    return None


def _looks_like_packing_token(tok: str) -> bool:
    t = str(tok or "").strip().rstrip(".,-_")
    if not t:
        return False
    if re.fullmatch(
        r"\d+(?:\.\d+)?(?:TAB|TABS|ML|MG|CAP|SYP|INJ|DS|S)", t, re.I
    ):
        return True
    if re.fullmatch(r"\d+\*\d+", t):
        return True
    # OCR packing like 10TAB_ / 5M. / 1OTAB (O for 0)
    if re.fullmatch(r"\d+(?:\.\d+)?(?:TAB|ML|MG|CAP|S)[A-Za-z._]*", t, re.I):
        return True
    if re.fullmatch(r"\d+[Oo]TAB[A-Za-z._]*", t, re.I):
        return True
    return False


def _repair_opbal_qty_tuple(use: List[float]) -> List[float]:
    """Rebuild Total as OpBal+Receipt when Total looks OCR-truncated."""
    if len(use) < 5:
        return use
    fixed = list(use)
    op, rec, total = fixed[0], fixed[1], fixed[2]
    expected = op + rec
    if abs(expected - total) <= max(1.0, 0.02 * max(abs(expected), 1.0)):
        return fixed
    if expected <= 0 or total < 0:
        return fixed
    exp_i, tot_i = int(round(expected)), int(round(total))
    if tot_i >= exp_i or tot_i <= 0:
        return fixed
    exp_s, tot_s = str(exp_i), str(tot_i)
    ratio = expected / max(total, 0.1)
    # 42→2 (suffix) or 60→6 (prefix / dropped trailing 0)
    if exp_s.endswith(tot_s) or (
        exp_s.startswith(tot_s) and (9 <= ratio <= 11 or 90 <= ratio <= 110)
    ):
        fixed[2] = float(expected)
    return fixed


def _parse_opbal_receipt_issue_row(
    ln: str, layout: str = "issue_closing_dump"
) -> Optional[Dict[str, Any]]:
    """Parse one OpBal/Receipt/Total/Issue/... product row."""
    if not ln or re.match(r"^\s*TOTAL\b", ln, re.I):
        return None
    if re.search(
        r"PRODUCT\s*NAME|PACKING|Op\.?\s*Bal|Page\s*No|Continued|Sales\s*&\s*Stock|"
        r"GRAND\s*TOTAL",
        ln,
        re.I,
    ):
        return None

    tokens = ln.split()
    if len(tokens) < 4:
        return None

    # Allow 8 tokens so a leading Shelf code can be dropped via last-7 slice
    qtys: List[float] = []
    i = len(tokens) - 1
    while i >= 0 and len(qtys) < 8:
        val = _ocr_qty_token(tokens[i])
        if val is None:
            break
        qtys.append(val)
        i -= 1
    qtys.reverse()
    if len(qtys) < 5:
        return None

    left = tokens[: i + 1]
    packing = None
    if left and _looks_like_packing_token(left[-1]):
        packing = left[-1].rstrip(".,-_")
        left = left[:-1]
    elif (
        len(left) >= 2
        and re.fullmatch(r"\d+(?:\.\d+)?", left[-2] or "")
        and re.fullmatch(r"(?:TAB|TABS|ML|MG|CAP|SYP|INJ|DS)", left[-1] or "", re.I)
    ):
        packing = f"{left[-2]}{left[-1]}"
        left = left[:-2]
    # Drop trailing bare shelf code left in the name (e.g. "... TAB 11")
    if left and re.fullmatch(r"\d{1,2}", left[-1] or ""):
        left = left[:-1]

    product_name = _clean_name(" ".join(left))
    if len(product_name) < 3:
        return None
    if re.fullmatch(r"AUROBINDO.*|VERITAZ.*|Qy\.?", product_name, re.I):
        return None

    # Prefer full 7-col layout, then 6, then 5. Repair truncated Total before accepting.
    use = None
    for n in (7, 6, 5):
        if n > len(qtys):
            continue
        candidate = _repair_opbal_qty_tuple(list(qtys[-n:] if len(qtys) > n else qtys))
        bal_diff = abs((candidate[0] + candidate[1]) - candidate[2])
        if bal_diff <= max(5.0, 0.25 * max(abs(candidate[2]), abs(candidate[0] + candidate[1]), 1.0)):
            use = candidate
            break
    if use is None:
        return None

    item = empty_line_item()
    item["product_name"] = product_name
    item["packing"] = packing
    item["opening_qty"] = use[0]
    item["receipts_qty"] = use[1]
    item["sales_qty"] = use[3]  # ISSUE (not Total)
    item["sales_value"] = 0.0
    item["closing_value"] = 0.0
    item["extra"]["total_stock_qty"] = use[2]

    if layout == "issue_expiry_closing" and len(use) >= 6:
        # OpBal Receipt Total Issue Expiry Closing [Near]
        item["extra"]["expiry_breakage_qty"] = use[4]
        item["closing_qty"] = use[5]
        if len(use) >= 7:
            item["extra"]["near_expiry_qty"] = use[6]
    else:
        # OpBal Receipt Total Issue Closing [Dump] [Near]
        item["closing_qty"] = use[4]
        if len(use) >= 6:
            item["extra"]["dump_qty"] = use[5]
        if len(use) >= 7:
            item["extra"]["near_expiry_qty"] = use[6]
    return item


def _parse_opbal_receipt_issue_statement(
    text: str, filename: str, source_format: str = "pdf"
) -> Optional[Dict[str, Any]]:
    """Parse Mahajan-style OpBal/Receipt/Total/Issue/Closing statements."""
    if not _is_opbal_issue_closing_format(text):
        return None

    layout = _opbal_layout(text)
    result = empty_result(filename, source_format)
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if lines:
        result["stockist_name"] = _clean_name(lines[0])
    for ln in lines[:20]:
        if re.search(r"Sales\s*&\s*Stock|Stock\s*&\s*Sales", ln, re.I):
            result["report_title"] = _clean_name(ln)
            m = re.search(
                r"(?:From\s+)?(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})\s*(?:Upto|to|[-–])\s*"
                r"(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})",
                ln,
                re.I,
            )
            if m:
                result["period_from"] = _normalize_date(m.group(1))
                result["period_to"] = _normalize_date(m.group(2))
        if re.search(r"AUROBINDO|VERITAZ|PHARMA\s+LIMITED|AUROBINO", ln, re.I) and not re.search(
            r"MAHAJAN|ASSOCIATES|Sales\s*&\s*Stock", ln, re.I
        ):
            result["company_name"] = _clean_name(ln)
        if re.search(r"GARH|JAMMU|MOHALLA|ROAD|NAGAR|R\.?\s*N\.?\s*PURA", ln, re.I) and not result.get(
            "stockist_address"
        ):
            result["stockist_address"] = _clean_name(ln)

    items: List[Dict[str, Any]] = []
    for ln in lines:
        row = _parse_opbal_receipt_issue_row(ln, layout=layout)
        if row:
            items.append(row)

    if not items:
        return None

    result["line_items"] = items
    result["totals"]["extra"]["extraction_method"] = "opbal_issue_closing_parser"
    result["totals"]["extra"]["opbal_layout"] = layout
    result["totals"]["sales_value"] = None
    result["totals"]["closing_value"] = None
    result = _apply_total_row_to_result(result, text)
    return result


def _parse_statement_total_row(text: str) -> Optional[Dict[str, Any]]:
    """Parse bottom TOTAL qty row (e.g. P.S.PHARMACEUTICALS OPENING/RECEIPT/ISSUE/CLOSING)."""
    if not text:
        return None

    header_blob = "\n".join(
        ln for ln in text.splitlines()[:50] if ln.strip()
    )
    # Prefer GRAND TOTAL, else last TOTAL line (footer)
    total_line = None
    grand_line = None
    for ln in text.splitlines():
        s = ln.strip()
        if re.match(r"^\s*GRAND\s*TOTAL\b", s, re.I):
            grand_line = s
        elif re.match(r"^\s*TOTAL\b", s, re.I):
            total_line = s
    total_line = grand_line or total_line

    if not total_line:
        return None

    # Strip currency/OCR junk before digit scan ($29770 → 29770; prefer clean GRAND)
    cleaned = total_line.replace("$", " ").replace(",", "")
    raw_nums = re.findall(r"-?\d+(?:\.\d+)?", cleaned)
    # Drop OCR-merged monsters (>8 integer digits)
    nums = [
        _to_float(n)
        for n in raw_nums
        if len(n.split(".")[0].lstrip("-")) <= 8
    ]
    if len(nums) < 4:
        return None

    parsed: Dict[str, Any] = {
        "total_row_raw": total_line,
        "total_row_numbers": nums,
    }

    # P.S.PHARMACEUTICALS: ITEM DESCRIPTION OPENING RECEIPT ISSUE CLOSING
    if re.search(r"OPENING\s+RECEIPT\s+ISSUE\s+CLOSING", header_blob, re.I) or (
        re.search(r"\bOPENING\b", header_blob, re.I)
        and re.search(r"\bISSUE\b", header_blob, re.I)
        and len(nums) == 4
    ):
        parsed.update(
            {
                "opening_qty": nums[0],
                "receipts_qty": nums[1],
                "sales_qty": nums[2],  # ISSUE
                "closing_qty": nums[3],
                "total_row_format": "opening_receipt_issue_closing",
                "total_row_labels": ["OPENING", "RECEIPT", "ISSUE", "CLOSING"],
            }
        )
        return parsed

    # Mahajan & Mahajan: OpBal Receipt Total Issue Expiry Closing Near
    if _opbal_layout(text) == "issue_expiry_closing" and len(nums) >= 6:
        parsed.update(
            {
                "opening_qty": nums[0],
                "receipts_qty": nums[1],
                "total_stock_qty": nums[2],
                "sales_qty": nums[3],  # ISSUE
                "expiry_breakage_qty": nums[4],
                "closing_qty": nums[5],  # Closing Balance (NOT the Expiry column)
                "total_row_format": "opbal_receipt_total_issue_expiry_closing",
                "total_row_labels": [
                    "OpBal",
                    "Receipt",
                    "Total",
                    "Issue",
                    "Expiry",
                    "Closing",
                ],
            }
        )
        if len(nums) >= 7:
            parsed["near_expiry_qty"] = nums[6]
            parsed["total_row_labels"].append("NearExpiry")
        return parsed

    # Mahajan Associates: OpBal Receipt Total Issue Closing [Dump] [Near Expiry]
    if re.search(r"Op\.?\s*Bal|Receipt|Issue|Closing", header_blob, re.I) and len(nums) >= 5:
        parsed.update(
            {
                "opening_qty": nums[0],
                "receipts_qty": nums[1],
                "total_stock_qty": nums[2],
                "sales_qty": nums[3],  # ISSUE
                "closing_qty": nums[4],
                "total_row_format": "opbal_receipt_total_issue_closing",
                "total_row_labels": [
                    "OpBal",
                    "Receipt",
                    "Total",
                    "Issue",
                    "Closing",
                ],
            }
        )
        if len(nums) >= 6:
            parsed["dump_qty"] = nums[5]
            parsed["total_row_labels"].append("Dump")
        if len(nums) >= 7:
            parsed["near_expiry_qty"] = nums[6]
            parsed["total_row_labels"].append("NearExpiry")
        return parsed

    # Generic 4-number TOTAL → treat as opening/receipt/issue/closing
    if len(nums) == 4:
        parsed.update(
            {
                "opening_qty": nums[0],
                "receipts_qty": nums[1],
                "sales_qty": nums[2],
                "closing_qty": nums[3],
                "total_row_format": "qty4_generic",
                "total_row_labels": ["OPENING", "RECEIPT", "ISSUE", "CLOSING"],
            }
        )
        return parsed

    return parsed


def _apply_total_row_to_result(result: Dict[str, Any], text: str) -> Dict[str, Any]:
    """Attach parsed footer TOTAL qty columns onto statement totals."""
    if not isinstance(result, dict) or result.get("multi_statement"):
        return result
    parsed = _parse_statement_total_row(text)
    if not parsed:
        return result

    totals = result.setdefault(
        "totals", {"sales_value": None, "closing_value": None, "extra": {}}
    )
    if not isinstance(totals.get("extra"), dict):
        totals["extra"] = {}

    # Authoritative qty totals from footer row (do not treat as money)
    for key in (
        "opening_qty",
        "receipts_qty",
        "sales_qty",
        "closing_qty",
        "total_stock_qty",
        "dump_qty",
        "near_expiry_qty",
        "expiry_breakage_qty",
    ):
        if key in parsed:
            totals[key] = parsed[key]

    totals["extra"]["total_row_raw"] = parsed.get("total_row_raw")
    totals["extra"]["total_row_numbers"] = parsed.get("total_row_numbers")
    totals["extra"]["total_row_format"] = parsed.get("total_row_format")
    totals["extra"]["total_row_labels"] = parsed.get("total_row_labels")
    totals["extra"]["total_row_source"] = "footer_total"

    # Qty-only formats must not keep footer numbers as sales_value money
    if parsed.get("total_row_format") in {
        "opening_receipt_issue_closing",
        "opbal_receipt_total_issue_closing",
        "opbal_receipt_total_issue_expiry_closing",
        "qty4_generic",
    }:
        if totals.get("sales_value") in parsed.get("total_row_numbers", []):
            totals["sales_value"] = None
        if totals.get("closing_value") in parsed.get("total_row_numbers", []):
            # closing_value money vs closing_qty — clear only if it matches a qty total
            totals["closing_value"] = None

    return result


def _is_implausible_money(value: Any, *, max_digits: int = 8) -> bool:
    """True for OCR-garbage currency (e.g. 75374123024 from merged TOTAL columns)."""
    if value is None:
        return False
    try:
        num = float(value)
    except (TypeError, ValueError):
        return True
    if num < 0:
        return True
    # Secondary-sales stockist totals almost never need > 8 integer digits
    int_digits = len(str(int(abs(num))))
    if int_digits > max_digits:
        return True
    # Hard cap (~10 crore) for a single stockist month statement value
    if abs(num) >= 100_000_000:
        return True
    return False


def _looks_like_concatenated_totals(sales_value: Any, closing_value: Any) -> bool:
    """Detect OCR merge of adjacent TOTAL numbers (75374|123024 -> 75374123024)."""
    try:
        sales = float(sales_value)
        closing = float(closing_value) if closing_value is not None else None
    except (TypeError, ValueError):
        return False
    if sales <= 0:
        return False
    sales_int = str(int(sales))
    if len(sales_int) < 9:
        return False
    if closing is not None and closing > 0:
        close_int = str(int(closing))
        # e.g. ...12302438126 where closing was partially merged — less common
        if sales_int.endswith(close_int) and len(sales_int) > len(close_int) + 3:
            return True
    # Split into plausible adjacent qty totals (4-6 digits + 4-6 digits)
    for split_at in range(4, min(7, len(sales_int) - 3)):
        left, right = sales_int[:split_at], sales_int[split_at:]
        if 3 <= len(right) <= 6 and left.isdigit() and right.isdigit():
            if int(left) > 0 and int(right) > 0:
                return True
    return False


def _sanitize_statement_financials(result: Dict[str, Any]) -> Dict[str, Any]:
    """Fix bogus sales/closing money totals from OCR TOTAL merges (qty-only formats)."""
    if not isinstance(result, dict):
        return result
    # Multi-statement wrapper
    if result.get("multi_statement") and isinstance(result.get("statements"), list):
        result["statements"] = [
            _sanitize_statement_financials(stmt) for stmt in result["statements"]
        ]
        return result

    items = result.get("line_items") or []
    if not isinstance(items, list):
        items = []

    sum_sales_value = sum(_to_float(i.get("sales_value")) for i in items if isinstance(i, dict))
    sum_closing_value = sum(_to_float(i.get("closing_value")) for i in items if isinstance(i, dict))
    sum_sales_qty = sum(_to_float(i.get("sales_qty")) for i in items if isinstance(i, dict))
    sum_closing_qty = sum(_to_float(i.get("closing_qty")) for i in items if isinstance(i, dict))

    totals = result.setdefault("totals", {"sales_value": None, "closing_value": None, "extra": {}})
    if not isinstance(totals.get("extra"), dict):
        totals["extra"] = {}

    totals["extra"]["line_sales_value_sum"] = sum_sales_value
    totals["extra"]["line_closing_value_sum"] = sum_closing_value

    # Prefer footer TOTAL-row qty when present (P.S.PHARMACEUTICALS etc.)
    has_footer_qty = totals.get("extra", {}).get("total_row_source") == "footer_total"
    if has_footer_qty:
        if totals.get("sales_qty") is not None:
            totals["extra"]["sales_qty"] = totals.get("sales_qty")
        else:
            totals["extra"]["sales_qty"] = sum_sales_qty
        if totals.get("closing_qty") is not None:
            totals["extra"]["closing_qty"] = totals.get("closing_qty")
        else:
            totals["extra"]["closing_qty"] = sum_closing_qty
    else:
        totals["extra"]["sales_qty"] = sum_sales_qty
        totals["extra"]["closing_qty"] = sum_closing_qty

    sales_total = totals.get("sales_value")
    closing_total = totals.get("closing_value")

    qty_only = (
        len(items) > 0
        and sum_sales_qty > 0
        and sum_sales_value <= 0
        and sum(1 for i in items if isinstance(i, dict) and _to_float(i.get("sales_value")) > 0)
        <= max(1, len(items) // 10)
    )

    bogus_sales = (
        _is_implausible_money(sales_total)
        or _looks_like_concatenated_totals(sales_total, closing_total)
        or (
            sales_total is not None
            and sum_sales_value > 0
            and float(sales_total) > max(sum_sales_value * 50, sum_sales_value + 100000)
        )
        or (qty_only and sales_total is not None and float(sales_total) > 0)
    )

    if bogus_sales:
        totals["extra"]["rejected_sales_value"] = sales_total
        # Prefer real line-item money sum; else null for qty-only statements
        totals["sales_value"] = sum_sales_value if sum_sales_value > 0 else None

    if _is_implausible_money(totals.get("closing_value")):
        totals["extra"]["rejected_closing_value"] = totals.get("closing_value")
        totals["closing_value"] = sum_closing_value if sum_closing_value > 0 else None
    elif (
        qty_only
        and totals.get("closing_value") is not None
        and sum_closing_value <= 0
        and float(totals.get("closing_value") or 0) > sum_closing_qty * 1000
        and float(totals.get("closing_value") or 0) > 100000
    ):
        # Closing "value" is likely a qty total mislabeled from OCR TOTAL
        totals["extra"]["rejected_closing_value"] = totals.get("closing_value")
        totals["extra"]["closing_qty_from_total_row"] = totals.get("closing_value")
        totals["closing_value"] = None

    # Sanitize absurd per-line sales_value (rare, but same OCR merge class)
    for item in items:
        if not isinstance(item, dict):
            continue
        if _is_implausible_money(item.get("sales_value")):
            item.setdefault("extra", {})
            if isinstance(item["extra"], dict):
                item["extra"]["rejected_sales_value"] = item.get("sales_value")
            item["sales_value"] = 0.0
        if _is_implausible_money(item.get("closing_value")):
            item.setdefault("extra", {})
            if isinstance(item["extra"], dict):
                item["extra"]["rejected_closing_value"] = item.get("closing_value")
            item["closing_value"] = 0.0

    return result


def _sniff_extension(file_bytes: bytes, filename: str) -> str:
    """Resolve extension from filename, with magic-byte fallback for images/text/Word."""
    name = Path(filename or "upload").name
    ext = Path(name).suffix.lower()
    if ext in SUPPORTED_EXTENSIONS:
        return ext

    head = file_bytes[:16] if file_bytes else b""
    if head.startswith(b"%PDF"):
        return ".pdf"
    if head.startswith(b"\xff\xd8\xff"):
        return ".jpg"
    if head.startswith(b"\x89PNG\r\n\x1a\n"):
        return ".png"
    if head.startswith(b"BM"):
        return ".bmp"
    if head.startswith(b"RIFF") and file_bytes[8:12] == b"WEBP":
        return ".webp"
    if head[:4] in (b"II*\x00", b"MM\x00*"):
        return ".tif"
    # OOXML Word (.docx) is a ZIP containing word/document.xml
    if head.startswith(b"PK"):
        try:
            import zipfile

            with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
                names = set(zf.namelist())
                if "word/document.xml" in names:
                    return ".docx"
                if "xl/workbook.xml" in names:
                    return ".xlsx"
        except Exception:
            pass
    # Legacy OLE: Word .doc vs Excel .xls
    if head.startswith(b"\xd0\xcf\x11\xe0"):
        try:
            import olefile

            with olefile.OleFileIO(io.BytesIO(file_bytes)) as ole:
                streams = {"/".join(p).lower() for p in ole.listdir()}
                if any("worddocument" in s for s in streams):
                    return ".doc"
                if any(s in streams for s in ("workbook", "book")):
                    return ".xls"
        except Exception:
            # Prefer .doc when filename hints Word; else leave as-is
            if "doc" in (ext or "") or "word" in name.lower():
                return ".doc"
    # UTF-8 / ASCII text sales statements
    if file_bytes and not any(b == 0 for b in file_bytes[:200]):
        try:
            sample = file_bytes[:2000].decode("utf-8")
            if re.search(r"STOCK|SALES|PRODUCT|OPENING|CLOSING", sample, re.I):
                return ".txt"
        except UnicodeDecodeError:
            try:
                sample = file_bytes[:2000].decode("latin-1")
                if re.search(r"STOCK|SALES|PRODUCT|OPENING|CLOSING", sample, re.I):
                    return ".txt"
            except Exception:
                pass
    return ext


def extract_sales_statement(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """Dispatch by extension and return unified sales-statement JSON."""
    name = Path(filename or "upload").name
    ext = _sniff_extension(file_bytes, name)
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported format '{ext or '(none)'}'. "
            f"Supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    if ext in TEXT_EXTENSIONS:
        result = _parse_txt(file_bytes, name)
    elif ext in {".htm", ".html"}:
        result = _parse_htm(file_bytes, name)
    elif ext in WORD_EXTENSIONS:
        result = _parse_word(file_bytes, name, ext)
    elif ext == ".pdf":
        result = _parse_pdf(file_bytes, name)
    elif ext in {".xls", ".xlsx"}:
        result = _parse_xls(file_bytes, name, ext)
    elif ext in IMAGE_EXTENSIONS:
        result = _parse_image(file_bytes, name, ext)
    else:
        raise ValueError(f"Unsupported format '{ext}'")

    return _sanitize_statement_financials(result)


# ---------------------------------------------------------------------------
# TXT (Kaveri-style fixed / spaced rows)
# ---------------------------------------------------------------------------

_KAVERI_ROW = re.compile(
    r"^(\d{3,5})\s+(.+?)\s+(\S*\*\S*|\S*(?:ML|MG|TAB|CAP|SYP|VIAL|DROPS|'S|S)\S*)\s+"
    r"(-?\d+(?:\.\d+)?)\s+(-?\d+(?:\.\d+)?)\s+(-?\d+(?:\.\d+)?)\s+"
    r"(-?\d+(?:\.\d+)?)\s+(-?\d+(?:\.\d+)?)\s+(-?\d+(?:\.\d+)?)\s+(-?\d+(?:\.\d+)?)"
    r"(?:\s+(\S+))?(?:\s+(\S+))?\s*$",
    re.IGNORECASE,
)


def _parse_txt(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    text = file_bytes.decode("utf-8", errors="ignore")
    if "\x00" in text[:200]:
        text = file_bytes.decode("latin-1", errors="ignore")

    result = empty_result(filename, "txt")
    lines = [ln.rstrip() for ln in text.splitlines()]

    # Header: stockist on first non-empty line
    non_empty = [ln.strip() for ln in lines if ln.strip() and not ln.strip().startswith("=")]
    if non_empty:
        result["stockist_name"] = _clean_name(non_empty[0])
    if len(non_empty) > 1 and not re.search(
        r"MONTHLY|STOCK|SALES|COMPANY|FROM", non_empty[1], re.I
    ):
        result["stockist_address"] = _clean_name(non_empty[1])

    for ln in lines:
        if re.search(r"MONTHLY\s+STOCK|STOCK\s*&\s*SALES|STOCK AND SALES", ln, re.I):
            result["report_title"] = _clean_name(ln)
        m_co = re.search(r"COMPANY\s*NAME\s*:?\s*(.+)$", ln, re.I)
        if m_co and m_co.group(1).strip():
            result["company_name"] = _clean_name(m_co.group(1))
        # COMPANY NAME on previous line, value on next — handled below
        m_from = re.search(
            r"(?:FROM|FORM)\s*:?\s*(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})\s*(?:TO|:)?\s*"
            r"(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})?",
            ln,
            re.I,
        )
        if m_from:
            result["period_from"] = _normalize_date(m_from.group(1))
            if m_from.group(2):
                result["period_to"] = _normalize_date(m_from.group(2))

    # Company name may be on its own line after "COMPANY NAME :"
    for i, ln in enumerate(lines):
        if re.search(r"COMPANY\s*NAME\s*:?\s*$", ln.strip(), re.I) and i + 1 < len(lines):
            nxt = lines[i + 1].strip()
            if nxt and not nxt.startswith("="):
                result["company_name"] = _clean_name(nxt)
                break
        if re.search(r"COMPANY\s*NAME\s*:", ln, re.I):
            after = re.split(r"COMPANY\s*NAME\s*:", ln, flags=re.I)[-1].strip()
            if after:
                result["company_name"] = _clean_name(after)

    items: List[Dict[str, Any]] = []
    for ln in lines:
        raw = ln.strip()
        if not raw or raw.startswith("=") or raw.startswith("PRD") or raw.startswith("CODE"):
            continue
        if re.search(r"^\s*TOTAL\s*:", raw, re.I):
            nums = re.findall(r"-?\d+(?:\.\d+)?", raw)
            if nums:
                # Product-grid total has 3 amounts (receipts/sales/closing).
                # Trailing voucher TOTAL with a single 0.00 must not overwrite.
                floats = [_to_float(n) for n in nums]
                if len(floats) >= 3:
                    result["totals"]["extra"]["receipts_value"] = floats[-3]
                    result["totals"]["sales_value"] = floats[-2]
                    result["totals"]["closing_value"] = floats[-1]
                elif len(floats) >= 2 and result["totals"]["sales_value"] is None:
                    result["totals"]["sales_value"] = floats[-2]
                    result["totals"]["closing_value"] = floats[-1]
            continue
        if re.match(r"^(VOTNO|Signature|FORM|FROM|TIME|DATE|PAGE|COMPANY)", raw, re.I):
            continue

        item = _parse_kaveri_line(raw)
        if item:
            items.append(item)

    result["line_items"] = items
    return result


def _parse_kaveri_line(line: str) -> Optional[Dict[str, Any]]:
    m = _KAVERI_ROW.match(line)
    if m:
        item = empty_line_item()
        item["product_code"] = m.group(1)
        item["product_name"] = _clean_name(m.group(2))
        item["packing"] = m.group(3)
        item["opening_qty"] = _to_float(m.group(4))
        item["receipts_qty"] = _to_float(m.group(5))
        item["extra"]["total_stock"] = _to_float(m.group(6))
        item["sales_qty"] = _to_float(m.group(7))
        item["sales_value"] = _to_float(m.group(8))
        item["closing_qty"] = _to_float(m.group(9))
        item["closing_value"] = _to_float(m.group(10))
        if m.group(11):
            item["extra"]["exp_batch"] = m.group(11)
        if m.group(12):
            item["extra"]["exp_date"] = m.group(12)
        return item

    # Fallback: code + right-aligned numeric fields
    m2 = re.match(r"^(\d{3,5})\s+(.+)$", line)
    if not m2:
        return None
    code = m2.group(1)
    rest = m2.group(2).strip()
    # Pull last 7–8 numbers (opening..closing_value [+ batch/date ignored if non-numeric])
    tokens = rest.split()
    if len(tokens) < 8:
        return None
    # Find packing token (contains * preferred)
    pack_idx = None
    for i, tok in enumerate(tokens):
        if "*" in tok:
            pack_idx = i
            break
    if pack_idx is None:
        # packing often just before opening qty sequence
        # take last 7 numeric tokens as metrics
        nums: List[str] = []
        cut = len(tokens)
        for i in range(len(tokens) - 1, -1, -1):
            if re.fullmatch(r"-?\d+(?:\.\d+)?", tokens[i]):
                nums.append(tokens[i])
                cut = i
                if len(nums) >= 7:
                    break
            elif nums:
                break
        nums = list(reversed(nums))
        if len(nums) < 7:
            return None
        name_pack = " ".join(tokens[:cut]).strip()
        packing = None
        product_name = name_pack
    else:
        product_name = " ".join(tokens[:pack_idx]).strip()
        packing = tokens[pack_idx]
        after = tokens[pack_idx + 1 :]
        nums = [t for t in after if re.fullmatch(r"-?\d+(?:\.\d+)?", t)]
        if len(nums) < 7:
            return None

    item = empty_line_item()
    item["product_code"] = code
    item["product_name"] = _clean_name(product_name)
    item["packing"] = packing
    item["opening_qty"] = _to_float(nums[0])
    item["receipts_qty"] = _to_float(nums[1])
    item["extra"]["total_stock"] = _to_float(nums[2])
    item["sales_qty"] = _to_float(nums[3])
    item["sales_value"] = _to_float(nums[4])
    item["closing_qty"] = _to_float(nums[5])
    item["closing_value"] = _to_float(nums[6])
    return item


# ---------------------------------------------------------------------------
# HTM (Livem-style)
# ---------------------------------------------------------------------------

def _map_livem_numbers(nums: List[float]) -> Dict[str, Any]:
    """Map Livem numeric cells to unified fields.

    Livem columns (when present):
      LMS (opening) | TotalStock Qty | Stock Amount | Sales Qty | Sales Amt | Close Qty | Close Amt

    The "Purchase Qty" cell is total available stock (opening + receipts), not receipts alone.
    """
    item = empty_line_item()
    if len(nums) >= 7:
        opening = nums[0]
        total_stock = nums[1]
        item["opening_qty"] = opening
        item["receipts_qty"] = max(0.0, total_stock - opening)
        item["extra"]["total_stock"] = total_stock
        item["extra"]["stock_value"] = nums[2]
        item["sales_qty"] = nums[3]
        item["sales_value"] = nums[4]
        item["closing_qty"] = nums[5]
        item["closing_value"] = nums[6]
    elif len(nums) == 6:
        opening = nums[0]
        total_stock = nums[1]
        item["opening_qty"] = opening
        item["receipts_qty"] = max(0.0, total_stock - opening)
        item["extra"]["total_stock"] = total_stock
        item["extra"]["stock_value"] = nums[2]
        item["sales_qty"] = nums[3]
        item["sales_value"] = nums[4]
        item["closing_qty"] = nums[5]
    elif len(nums) == 5:
        # LMS, TotalStock, Amount, Close Qty, Close Amt (no sales movement)
        opening = nums[0]
        total_stock = nums[1]
        item["opening_qty"] = opening
        item["receipts_qty"] = max(0.0, total_stock - opening)
        item["extra"]["total_stock"] = total_stock
        item["extra"]["stock_value"] = nums[2]
        item["closing_qty"] = nums[3]
        item["closing_value"] = nums[4]
    elif len(nums) == 4:
        # LMS, Opening/Close value, Close Qty, Close Amt — no purchase/sales
        item["opening_qty"] = nums[0]
        item["extra"]["opening_value"] = nums[1]
        item["closing_qty"] = nums[2]
        item["closing_value"] = nums[3]
        if abs(nums[0] - nums[2]) < 1e-9 and abs(nums[1] - nums[3]) < 1e-9:
            # mirrored open/close — treat amount as both opening and closing value
            pass
    elif len(nums) == 3:
        item["opening_qty"] = nums[0]
        item["closing_qty"] = nums[1]
        item["closing_value"] = nums[2]
    elif len(nums) == 2:
        item["closing_qty"] = nums[0]
        item["closing_value"] = nums[1]
        item["opening_qty"] = nums[0]
    elif len(nums) == 1:
        item["closing_qty"] = nums[0]
        item["opening_qty"] = nums[0]
    else:
        item["extra"]["raw_numbers"] = nums
    return item


def _parse_htm(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    from bs4 import BeautifulSoup

    result = empty_result(filename, "htm")
    html = file_bytes.decode("latin-1", errors="ignore")
    soup = BeautifulSoup(html, "html.parser")

    rows: List[List[str]] = []
    for tr in soup.find_all("tr"):
        cells = [c.get_text(" ", strip=True) for c in tr.find_all(["td", "th"])]
        cells = [c for c in cells if c]
        if cells:
            rows.append(cells)

    for cells in rows:
        joined = " ".join(cells)
        if len(cells) == 1 and not re.search(r"_+|Company|Product|TOTAL|SUB", cells[0], re.I):
            if not result["stockist_name"]:
                result["stockist_name"] = _clean_name(cells[0])
        if "Company" in joined:
            for i, c in enumerate(cells):
                if re.search(r"Company\s*:", c, re.I) and i + 1 < len(cells):
                    result["company_name"] = _clean_name(cells[i + 1])
                elif re.search(r"Company\s*:\s*(.+)", c, re.I):
                    result["company_name"] = _clean_name(
                        re.search(r"Company\s*:\s*(.+)", c, re.I).group(1)
                    )
            m_title = re.search(r"Stock and Sales Report.*", joined, re.I)
            if m_title:
                result["report_title"] = _clean_name(m_title.group(0))
            m_mo = re.search(r"Month\s+Of\s+([A-Za-z]+)\s+(\d{4})", joined, re.I)
            if m_mo:
                pf, pt = _month_period(m_mo.group(1), int(m_mo.group(2)))
                result["period_from"], result["period_to"] = pf, pt

    items: List[Dict[str, Any]] = []
    for cells in rows:
        if len(cells) < 2:
            continue
        name = cells[0]
        if re.search(
            r"^(Product|Opening|Purchase|Sale|Closing|Company|SUB|Total|_+|LIVEM)",
            name,
            re.I,
        ):
            if re.search(r"^SUB\s*TOTAL|^Total\s*:", name, re.I) or (
                len(cells) >= 2 and re.search(r"^Total\s*:", name, re.I)
            ):
                nums = [_to_float(c) for c in cells if re.search(r"\d", c)]
                # opening_amt, purchase_return?, sales, closing
                if len(nums) >= 4:
                    result["totals"]["extra"]["opening_value"] = nums[-4]
                    result["totals"]["extra"]["purchase_return_value"] = nums[-3]
                    result["totals"]["sales_value"] = nums[-2]
                    result["totals"]["closing_value"] = nums[-1]
                elif len(nums) >= 2:
                    result["totals"]["sales_value"] = nums[-2]
                    result["totals"]["closing_value"] = nums[-1]
            continue
        if re.fullmatch(r"_+", name) or name.startswith("_"):
            continue

        packing = cells[1] if len(cells) > 1 else None
        nums_start = 2
        if packing and re.fullmatch(r"-?\d+(?:\.\d+)?", packing.replace(",", "")):
            packing = None
            nums_start = 1

        num_cells = cells[nums_start:]
        nums = [_to_float(c) for c in num_cells if re.search(r"\d", c)]

        # Source rows that only list product + pack have no stock figures in the HTM.
        # Skip them so the response is not filled with all-zero placeholders.
        if not nums:
            continue

        item = _map_livem_numbers(nums)
        item["product_name"] = _clean_name(name)
        item["packing"] = packing
        items.append(item)

    result["line_items"] = items
    return result


# ---------------------------------------------------------------------------
# Word (.doc / .docx)
# ---------------------------------------------------------------------------

def _parse_word(file_bytes: bytes, filename: str, ext: str) -> Dict[str, Any]:
    """Parse Word sales statements (.docx via python-docx; .doc via convert/fallback)."""
    # Misnamed or sniffed OOXML
    if ext == ".docx" or file_bytes[:2] == b"PK":
        return _parse_docx(file_bytes, filename)

    # Legacy .doc → try convert to docx, else extract text
    converted = _convert_doc_to_docx_bytes(file_bytes)
    if converted:
        result = _parse_docx(converted, filename)
        result["source_format"] = "doc"
        return result

    text = _extract_legacy_doc_text(file_bytes)
    if text.strip():
        # Reuse TXT parser heuristics on extracted Word text
        result = _parse_txt(text.encode("utf-8", errors="ignore"), filename)
        result["source_format"] = "doc"
        result["totals"]["extra"]["extraction_method"] = "legacy_doc_text"
        return result

    raise ValueError(
        "Could not parse legacy .doc file. Please re-save as .docx and upload again."
    )


def _convert_doc_to_docx_bytes(file_bytes: bytes) -> Optional[bytes]:
    """Convert legacy .doc to .docx via Microsoft Word COM when available (Windows)."""
    try:
        import pythoncom  # type: ignore
        import win32com.client  # type: ignore
    except ImportError:
        return None

    import os
    import tempfile

    in_path = out_path = None
    word = None
    try:
        pythoncom.CoInitialize()
        fd_in, in_path = tempfile.mkstemp(suffix=".doc")
        os.close(fd_in)
        fd_out, out_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd_out)
        with open(in_path, "wb") as fh:
            fh.write(file_bytes)

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(in_path, ReadOnly=True)
        # 16 = wdFormatXMLDocument (.docx)
        doc.SaveAs(out_path, FileFormat=16)
        doc.Close(False)
        with open(out_path, "rb") as fh:
            return fh.read()
    except Exception as exc:
        logger.warning("Word COM .doc conversion failed: %s", exc)
        return None
    finally:
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass
        for path in (in_path, out_path):
            if path and os.path.exists(path):
                try:
                    os.remove(path)
                except OSError:
                    pass


def _extract_legacy_doc_text(file_bytes: bytes) -> str:
    """Best-effort text extraction from legacy OLE .doc without Word installed."""
    chunks: List[str] = []
    try:
        import olefile

        with olefile.OleFileIO(io.BytesIO(file_bytes)) as ole:
            for stream_name in ("WordDocument", "1Table", "0Table"):
                if not ole.exists(stream_name):
                    continue
                raw = ole.openstream(stream_name).read()
                # Prefer UTF-16LE runs (common in Word binary)
                try:
                    decoded = raw.decode("utf-16-le", errors="ignore")
                    cleaned = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", decoded)
                    if len(cleaned.strip()) > 40:
                        chunks.append(cleaned)
                except Exception:
                    pass
                ascii_parts = re.findall(rb"[\x20-\x7e]{4,}", raw)
                if ascii_parts:
                    chunks.append(b" ".join(ascii_parts).decode("ascii", errors="ignore"))
    except Exception as exc:
        logger.warning("OLE .doc text extract failed: %s", exc)

    if not chunks:
        # Last resort: scan whole file for readable ASCII
        ascii_parts = re.findall(rb"[\x20-\x7e]{4,}", file_bytes)
        if ascii_parts:
            chunks.append(b"\n".join(ascii_parts).decode("ascii", errors="ignore"))

    return "\n".join(chunks)


def _norm_docx_header(header: Any) -> str:
    """Normalize Word table headers: 'Cr.Sch.Qty' / 'Op. Val' -> 'crschqty' / 'opval'."""
    return re.sub(r"[^a-z0-9]", "", str(header or "").strip().lower())


def _unglue_stock_sales_label(text: str) -> str:
    """Insert spaces into glued Stock-and-Sales Detail Report labels."""
    t = _clean_name(text)
    if not t:
        return t
    replacements = (
        (r"(?i)stockandsalesdetail", "Stock and Sales Detail"),
        (r"(?i)stockandsales", "Stock and Sales"),
        (r"(?i)\baurobindo(?=pharma)", "AUROBINDO "),
        (r"(?i)pharma(?=ltd|limited)", "PHARMA "),
        (r"(?i)\brajesh(?=medicos)", "RAJESH "),
        (r"(?i)\bmedicos\s*-\s*", "MEDICOS - "),
    )
    for pat, repl in replacements:
        t = re.sub(pat, repl, t)
    return _clean_name(t)


# Canonical Stock-and-Sales Detail Report columns (20-col Aurobindo / Medicos format)
_STOCK_SALES_DETAIL_HEADER_ALIASES: Dict[str, Tuple[str, ...]] = {
    "sl_no": ("slno", "sno", "srno"),
    "product_name": ("item", "product", "productname", "particulars"),
    "opening_qty": ("opqty", "openingqty", "opbal", "opbalqty"),
    "opening_value": ("opval", "openingval", "openingvalue", "opvalue"),
    "receipts_qty": ("pqty", "purchaseqty", "purqty", "receiptsqty"),
    "purchase_scheme": ("psch", "purchasescheme", "pursch", "purchasesch"),
    "purchase_value": ("pval", "purchaseval", "purchasevalue", "purval"),
    "sales_qty": ("sqty", "saleqty", "salesqty"),
    "sales_scheme": ("ssch", "salescheme", "salesscheme", "salesch"),
    "sales_value": ("sval", "saleval", "salesval", "salesvalue"),
    "break_sales_qty": ("brsqty", "brsaleqty", "breaksqty", "breakagesqty"),
    "break_sales_value": ("brsval", "brsaleval", "breaksval", "breakagesval"),
    "credit_qty": ("crqty", "creditqty", "crnoteqty"),
    "credit_scheme_qty": ("crschqty", "creditschemeqty", "crsch"),
    "credit_value": ("crval", "creditval", "creditvalue"),
    "debit_qty": ("dbqty", "debitqty", "drqty"),
    "debit_scheme_qty": ("dbschqty", "debitschemeqty", "dbsch", "drschqty"),
    "debit_value": ("dbval", "debitval", "debitvalue", "drval"),
    "closing_qty": ("clqty", "closingqty", "clqnt", "balqty", "closingbal"),
    "closing_value": ("clval", "closingval", "closingvalue", "clvalue", "balval"),
}


def _map_stock_sales_detail_headers(headers: List[Any]) -> Dict[str, int]:
    """Map normalized header cells onto canonical field names."""
    mapping: Dict[str, int] = {}
    for idx, header in enumerate(headers):
        key = _norm_docx_header(header)
        if not key:
            continue
        for field, aliases in _STOCK_SALES_DETAIL_HEADER_ALIASES.items():
            if field in mapping:
                continue
            if key == field.replace("_", "") or key in aliases:
                mapping[field] = idx
                break
    return mapping


def _cell_at(cells: List[str], idx: Optional[int]) -> str:
    if idx is None or idx < 0 or idx >= len(cells):
        return ""
    return str(cells[idx] if cells[idx] is not None else "").strip()


def _parse_stock_sales_detail_total_row(
    cells: List[str], colmap: Dict[str, int], result: Dict[str, Any]
) -> None:
    """Capture Total-row values so closing/sales match the printed report."""
    def _val(field: str, fallback_idx: Optional[int] = None) -> Optional[float]:
        raw = _cell_at(cells, colmap.get(field, fallback_idx))
        return _to_nullable_float(raw) if raw not in (None, "") else None

    sales = _val("sales_value", 9)
    closing = _val("closing_value", 19)
    if sales is not None:
        result["totals"]["sales_value"] = sales
    if closing is not None:
        result["totals"]["closing_value"] = closing

    extra = result["totals"].setdefault("extra", {})
    extra["raw_total_row"] = cells
    for field, fallback in (
        ("opening_value", 3),
        ("purchase_value", 6),
        ("break_sales_value", 11),
        ("credit_value", 14),
        ("debit_value", 17),
        ("opening_qty", 2),
        ("receipts_qty", 4),
        ("purchase_scheme", 5),
        ("sales_qty", 7),
        ("sales_scheme", 8),
        ("break_sales_qty", 10),
        ("credit_qty", 12),
        ("credit_scheme_qty", 13),
        ("debit_qty", 15),
        ("debit_scheme_qty", 16),
        ("closing_qty", 18),
    ):
        value = _val(field, fallback)
        if value is not None:
            extra[field] = value


def _parse_stock_sales_detail_item_row(
    cells: List[str], colmap: Dict[str, int]
) -> Optional[Dict[str, Any]]:
    """Build one line item with every Stock-and-Sales Detail column preserved."""
    # Prefer header map; fall back to fixed 20-col positions used by this report.
    positional = {
        "sl_no": 0,
        "product_name": 1,
        "opening_qty": 2,
        "opening_value": 3,
        "receipts_qty": 4,
        "purchase_scheme": 5,
        "purchase_value": 6,
        "sales_qty": 7,
        "sales_scheme": 8,
        "sales_value": 9,
        "break_sales_qty": 10,
        "break_sales_value": 11,
        "credit_qty": 12,
        "credit_scheme_qty": 13,
        "credit_value": 14,
        "debit_qty": 15,
        "debit_scheme_qty": 16,
        "debit_value": 17,
        "closing_qty": 18,
        "closing_value": 19,
    }

    def raw(field: str) -> str:
        return _cell_at(cells, colmap.get(field, positional.get(field)))

    product_name = _clean_name(raw("product_name").replace("\n", " ").replace("\r", " "))
    sl_no = raw("sl_no")
    if not product_name:
        return None
    # Skip repeated header rows inside later tables
    if _norm_docx_header(product_name) in {"item", "product"}:
        return None
    if sl_no and not re.match(r"^\d+$", sl_no) and _norm_docx_header(sl_no) in {
        "slno", "total", "sno"
    }:
        return None

    item = empty_line_item()
    item["product_name"] = product_name
    item["opening_qty"] = _to_float(raw("opening_qty"))
    item["receipts_qty"] = _to_float(raw("receipts_qty"))
    item["sales_qty"] = _to_float(raw("sales_qty"))
    item["sales_value"] = _to_float(raw("sales_value"))
    item["closing_qty"] = _to_float(raw("closing_qty"))
    item["closing_value"] = _to_float(raw("closing_value"))

    # Keep every remaining numeric column (including zeros) for reconciliation.
    extra_fields = (
        "opening_value",
        "purchase_scheme",
        "purchase_value",
        "sales_scheme",
        "break_sales_qty",
        "break_sales_value",
        "credit_qty",
        "credit_scheme_qty",
        "credit_value",
        "debit_qty",
        "debit_scheme_qty",
        "debit_value",
    )
    extra: Dict[str, Any] = {}
    if sl_no:
        extra["sl_no"] = sl_no
    for field in extra_fields:
        cell = raw(field)
        # Always store when the column exists in the sheet (header or positional)
        if field in colmap or len(cells) > positional[field]:
            extra[field] = _to_float(cell) if cell != "" else 0.0
    item["extra"] = extra
    return item


def _parse_docx(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    from docx import Document

    result = empty_result(filename, "docx")
    doc = Document(io.BytesIO(file_bytes))

    for para in doc.paragraphs:
        t = _clean_name(para.text)
        if not t:
            continue
        if re.search(r"Stock\s*and\s*Sales|StockandSales", t, re.I):
            result["report_title"] = _unglue_stock_sales_label(t)
        m_seller = re.search(r"Seller\s*:\s*(.+?)(?:\s+From|\s*$)", t, re.I)
        if m_seller:
            result["stockist_name"] = _unglue_stock_sales_label(m_seller.group(1))
        m_from = re.search(
            r"From\s*(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})\s*to\s*(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})",
            t,
            re.I,
        )
        if m_from:
            result["period_from"] = _normalize_date(m_from.group(1))
            result["period_to"] = _normalize_date(m_from.group(2))
        if "ReportDate" in t.replace(" ", "") or re.search(r"Report\s*Date", t, re.I):
            company = re.split(r"Report\s*Date", t, flags=re.I)[0]
            company = re.sub(r"\s+", " ", company).strip(" -\t")
            if company:
                result["company_name"] = _unglue_stock_sales_label(company)

    items: List[Dict[str, Any]] = []
    is_stock_sales_detail = False

    for table in doc.tables:
        if not table.rows:
            continue
        headers = [c.text.strip() for c in table.rows[0].cells]
        colmap = _map_stock_sales_detail_headers(headers)
        header_norm = {_norm_docx_header(h) for h in headers}
        table_is_detail = (
            "item" in header_norm
            and ("clval" in header_norm or "closingvalue" in header_norm)
            and ("opqty" in header_norm or "openingqty" in header_norm)
        )
        if table_is_detail:
            is_stock_sales_detail = True

        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if not cells or not any(cells):
                continue

            first_norm = _norm_docx_header(cells[0])
            if first_norm in {"slno", "sno", "srno"}:
                continue
            if first_norm == "total" or str(cells[0]).strip().lower() == "total":
                if table_is_detail or len(cells) >= 20:
                    _parse_stock_sales_detail_total_row(cells, colmap, result)
                else:
                    # Legacy narrow table fallback
                    if len(cells) > 9:
                        result["totals"]["sales_value"] = _to_nullable_float(cells[9])
                    if len(cells) > 19:
                        result["totals"]["closing_value"] = _to_nullable_float(cells[19])
                    result["totals"]["extra"]["raw"] = cells
                continue

            if table_is_detail or (
                "item" in header_norm and ("clval" in colmap or len(cells) >= 20)
            ):
                item = _parse_stock_sales_detail_item_row(cells, colmap)
                if item:
                    items.append(item)
                continue

            # Non Stock-and-Sales-Detail Word tables: keep previous best-effort path
            if not re.match(r"^\d+$", cells[0].strip()) and len(cells) < 3:
                continue
            item = empty_line_item()
            if len(cells) < 10:
                continue
            item["product_name"] = _clean_name(cells[1].replace("\n", " "))
            item["opening_qty"] = _to_float(cells[2])
            item["receipts_qty"] = _to_float(cells[4] if len(cells) > 4 else 0)
            item["sales_qty"] = _to_float(cells[7] if len(cells) > 7 else 0)
            item["sales_value"] = _to_float(cells[9] if len(cells) > 9 else 0)
            item["closing_qty"] = _to_float(cells[18] if len(cells) > 18 else 0)
            item["closing_value"] = _to_float(cells[19] if len(cells) > 19 else 0)
            item["extra"] = {
                "sl_no": cells[0],
                "opening_value": _to_float(cells[3] if len(cells) > 3 else 0),
                "purchase_scheme": _to_float(cells[5] if len(cells) > 5 else 0),
                "purchase_value": _to_float(cells[6] if len(cells) > 6 else 0),
                "sales_scheme": _to_float(cells[8] if len(cells) > 8 else 0),
            }
            if item["product_name"]:
                items.append(item)

    result["line_items"] = items
    if is_stock_sales_detail:
        result["totals"].setdefault("extra", {})["extraction_method"] = (
            "docx_stock_sales_detail"
        )
        # Prefer printed Total-row closing; if absent, fall back to line sum.
        if result["totals"].get("closing_value") is None and items:
            result["totals"]["closing_value"] = round(
                sum(_to_float(i.get("closing_value")) for i in items), 2
            )
        if result["totals"].get("sales_value") is None and items:
            result["totals"]["sales_value"] = round(
                sum(_to_float(i.get("sales_value")) for i in items), 2
            )
        line_closing = round(sum(_to_float(i.get("closing_value")) for i in items), 2)
        printed_closing = result["totals"].get("closing_value")
        result["totals"]["extra"]["line_closing_value_sum"] = line_closing
        if printed_closing is not None:
            result["totals"]["extra"]["closing_value_matches_lines"] = (
                abs(float(printed_closing) - line_closing) < 0.05
            )
    return result


# ---------------------------------------------------------------------------
# XLS / XLSX (Victory-style sparse grid)
# ---------------------------------------------------------------------------

def _parse_xls(file_bytes: bytes, filename: str, ext: str) -> Dict[str, Any]:
    result = empty_result(filename, ext.lstrip("."))

    rows: List[List[Any]] = []
    if ext == ".xls":
        import xlrd

        wb = xlrd.open_workbook(file_contents=file_bytes)
        sh = wb.sheet_by_index(0)
        for r in range(sh.nrows):
            rows.append([sh.cell_value(r, c) for c in range(sh.ncols)])
    else:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
        sh = wb.active
        for row in sh.iter_rows(values_only=True):
            rows.append(list(row))

    # Metadata scan
    for row in rows[:20]:
        texts = [str(c).strip() for c in row if c is not None and str(c).strip()]
        joined = " ".join(texts)
        if not texts:
            continue
        if not result["stockist_name"] and re.search(
            r"MEDICAL|STORES|AGENC|PHARMA|MEDICO", texts[0], re.I
        ):
            result["stockist_name"] = _clean_name(texts[0])
        if re.search(r"Phone|NAGAR|ROAD|HOUSE|COMPLEX", joined, re.I) and len(joined) > 20:
            if not result["stockist_address"]:
                result["stockist_address"] = _clean_name(joined)
        if re.search(r"Stock and Sale|Stock & Sale|Sales Report", joined, re.I):
            result["report_title"] = _clean_name(joined)
            m = re.search(
                r"From\s*date\s*(\d{1,2}[- ][A-Za-z]{3,9}[- ]\d{2,4})\s*to\s*"
                r"(\d{1,2}[- ][A-Za-z]{3,9}[- ]\d{2,4})",
                joined,
                re.I,
            )
            if m:
                result["period_from"] = _normalize_date(m.group(1))
                result["period_to"] = _normalize_date(m.group(2))
        if re.search(r"Manufacturer", joined, re.I):
            # e.g. Manufacturer 000255 AUROBINDO
            parts = [t for t in texts if not re.fullmatch(r"\d+", t)]
            for i, t in enumerate(texts):
                if re.search(r"Manufacturer", t, re.I) and i + 1 < len(texts):
                    # skip code-like next token
                    cand = texts[i + 1]
                    if re.fullmatch(r"\d+", str(cand)) and i + 2 < len(texts):
                        result["company_name"] = _clean_name(str(texts[i + 2]))
                    else:
                        result["company_name"] = _clean_name(str(cand))

    # Detect header row with Op. / Sale / Bal.
    header_idx = None
    colmap: Dict[str, int] = {}
    for ri, row in enumerate(rows):
        labels = {
            str(c).strip().lower(): ci
            for ci, c in enumerate(row)
            if c is not None and str(c).strip()
        }
        if "item" in labels and ("op." in labels or "op" in labels or "bal." in labels):
            header_idx = ri
            # normalize keys
            for lab, ci in labels.items():
                key = lab.replace(" ", "")
                colmap[key] = ci
            break

    def _col(*names: str) -> Optional[int]:
        for n in names:
            k = n.lower().replace(" ", "")
            if k in colmap:
                return colmap[k]
        return None

    items: List[Dict[str, Any]] = []
    start = (header_idx + 1) if header_idx is not None else 0
    for row in rows[start:]:
        texts = [(ci, c) for ci, c in enumerate(row) if c is not None and str(c).strip() != ""]
        if not texts:
            continue
        joined = " ".join(str(c) for _, c in texts)
        if re.search(
            r"For Manufacturer|Opening Val|Closing Val|(?<!\()\bSales\s*:|Printed Using|Report Date|Page \d",
            joined,
            re.I,
        ):
            if re.search(r"Opening Val", joined, re.I):
                for i, (_ci, c) in enumerate(texts):
                    if re.search(r"Opening Val", str(c), re.I) and i + 1 < len(texts):
                        result["totals"]["extra"]["opening_value"] = _to_float(texts[i + 1][1])
                    if re.search(r"Purchase", str(c), re.I) and i + 1 < len(texts):
                        result["totals"]["extra"]["purchase_value"] = _to_float(texts[i + 1][1])
            if re.search(r"Closing Val", joined, re.I):
                for i, (_ci, c) in enumerate(texts):
                    if re.search(r"Closing Val", str(c), re.I) and i + 1 < len(texts):
                        result["totals"]["closing_value"] = _to_float(texts[i + 1][1])
            if re.search(r"\bSales\s*:", joined, re.I):
                for i, (_ci, c) in enumerate(texts):
                    label = str(c).strip()
                    # Prefer primary "Sales :" over "Sales (May):" / "Sales (Apr):"
                    if re.fullmatch(r"Sales\s*:", label, re.I) and i + 1 < len(texts):
                        result["totals"]["sales_value"] = _to_float(texts[i + 1][1])
                        break
            continue
        if re.search(r"Manufacturer|Item\b|Pack\b", joined, re.I) and header_idx is not None:
            continue

        # Product row: name usually col 1
        name_idx = _col("item")
        product_name = None
        if name_idx is not None and name_idx < len(row) and row[name_idx]:
            product_name = str(row[name_idx]).strip()
        else:
            # first long text cell
            for ci, c in texts:
                s = str(c).strip()
                if len(s) > 3 and not re.fullmatch(r"-?\d+(?:\.\d+)?", s):
                    if not re.search(r"Manufacturer|AUROBINDO|Pack|Item", s, re.I):
                        product_name = s
                        break
        if not product_name:
            continue
        if re.search(r"^(Item|Pack|Manufacturer|Sales\s*:|Opening|Closing)$", product_name, re.I):
            continue
        if re.search(r"^(Sales|Opening Val|Closing Val|Purchase|Credit|Branch|Adj)\b", product_name, re.I):
            continue

        item = empty_line_item()
        item["product_name"] = _clean_name(product_name)

        pack_i = _col("pack")
        if pack_i is not None and pack_i < len(row) and row[pack_i]:
            item["packing"] = str(row[pack_i]).strip()

        op_i = _col("op.", "op")
        pur_i = _col("pur")
        sale_i = _col("sale")
        bal_i = _col("bal.", "bal")
        bval_i = _col("bval")
        sval_i = _col("sval")

        if op_i is not None:
            item["opening_qty"] = _to_float(row[op_i] if op_i < len(row) else 0)
        if pur_i is not None:
            item["receipts_qty"] = _to_float(row[pur_i] if pur_i < len(row) else 0)
        if sale_i is not None:
            item["sales_qty"] = _to_float(row[sale_i] if sale_i < len(row) else 0)
        if bal_i is not None:
            item["closing_qty"] = _to_float(row[bal_i] if bal_i < len(row) else 0)
        if bval_i is not None:
            item["closing_value"] = _to_float(row[bval_i] if bval_i < len(row) else 0)
        if sval_i is not None:
            item["sales_value"] = _to_float(row[sval_i] if sval_i < len(row) else 0)

        # May/Apr history columns if present
        for hist in ("may", "apr"):
            hi = _col(hist)
            if hi is not None and hi < len(row) and row[hi] not in (None, ""):
                item["extra"][f"hist_{hist}"] = _to_float(row[hi])

        items.append(item)

    result["line_items"] = items
    return result


# ---------------------------------------------------------------------------
# PDF (split multi-stockist statements, then extract each)
# ---------------------------------------------------------------------------

_STOCKIST_NAME_HINT = re.compile(
    r"PHARMACEUTICAL|PHARMAC(?:Y|IES)|AGENC|MEDICOSE|MEDICOS|"
    r"MEDICAL\s+(?:STORES|AGENC)|ASSOCIATES|DISTRIBUT|ENTERPRISES|"
    r"TRADERS?|AGENCIES",
    re.I,
)
_STATEMENT_TITLE_HINT = re.compile(
    r"STOCK\s*&\s*SALES|STOCK\s+AND\s+SALES|SALES\s*&\s*STOCK|"
    r"Sales\s*&\s*Stock\s+Statement|STOCK\s+REPORT|Stock\s+and\s+Sale",
    re.I,
)
_CONTINUATION_HINT = re.compile(
    r"^\s*(?:--\s*)?Continued\s+Page|^\s*Page\s*No\.?\s*[2-9]\b",
    re.I,
)
_MANUFACTURER_HINT = re.compile(
    r"\b(?:AUROBINDO|VERITAZ|HEALTHCARE\s+LTD|PHARMA\s+LTD|PHARMA\s+LIMITED|"
    r"LABORATOR)\b",
    re.I,
)


def _normalize_stockist_key(name: str) -> str:
    text = re.sub(r"\s+", " ", (name or "").upper()).strip()
    text = text.replace("&", " AND ")
    text = re.sub(r"[^A-Z0-9 ]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    # Drop trailing year/noise like "26-27"
    text = re.sub(r"\b\d{2}-\d{2}\b", "", text).strip()
    return text


def _looks_like_stockist_header(line: str) -> bool:
    s = (line or "").strip()
    if len(s) < 4 or len(s) > 90:
        return False
    if re.search(r"^\d", s):
        return False
    if re.search(
        r"PRODUCT|PACKING|ITEM\s+DESCRIPT|OpBal|OPENING|CLOSING|GSTIN|Phone\s*:|"
        r"E-Mail|Page\s*No|Continued|STOCK\s*&\s*SALES|Sales\s*&\s*Stock",
        s,
        re.I,
    ):
        return False
    if _MANUFACTURER_HINT.search(s) and not _STOCKIST_NAME_HINT.search(s):
        return False
    if _STOCKIST_NAME_HINT.search(s):
        return True
    # ALL-CAPS agency-like short header
    letters = re.sub(r"[^A-Za-z]", "", s)
    if letters and letters.isupper() and len(s.split()) <= 8:
        return True
    return False


def _clean_stockist_label(name: str) -> str:
    text = _clean_name(name)
    text = re.sub(r"\s+\d{2}-\d{2}\s*$", "", text).strip()
    return text


def _detect_stockist_from_page_text(text: str) -> Optional[str]:
    """Return stockist name if this page starts a (new) statement."""
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
    if not lines:
        return None

    # Continuations of previous stockist's statement
    head = "\n".join(lines[:6])
    if _CONTINUATION_HINT.search(lines[0]) and not _looks_like_stockist_header(lines[0]):
        return None
    if re.search(r"Continued\s+Page", head, re.I) and not _looks_like_stockist_header(
        lines[0]
    ):
        # Still allow if first line is clearly a new stockist
        if not _looks_like_stockist_header(lines[0]):
            return None

    # Prefer first stockist-like line near top
    for ln in lines[:8]:
        if _looks_like_stockist_header(ln):
            # Require statement context somewhere on page when possible
            if _STATEMENT_TITLE_HINT.search(text) or _STOCKIST_NAME_HINT.search(ln):
                return _clean_stockist_label(ln)
            return _clean_stockist_label(ln)
    return None


def _ocr_pdf_page_text(page, zoom: float = 2.0) -> Tuple[str, bytes]:
    """Render PDF page to PNG and OCR; return (text, png_bytes)."""
    import fitz

    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    img_bytes = pix.tobytes("png")
    try:
        text = _ocr_image_to_text(img_bytes)
    except Exception as exc:
        logger.warning("PDF page OCR failed: %s", exc)
        text = ""
    text = _merge_footer_total_into_text(text or "", img_bytes)
    return text or "", img_bytes


def _group_pdf_pages_by_stockist(
    page_infos: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    """
    Group page infos into stockist statement segments.

    page_infos items: {page_index, text, image_bytes, stockist}
    """
    groups: List[Dict[str, Any]] = []
    for info in page_infos:
        stockist = info.get("stockist")
        if stockist:
            key = _normalize_stockist_key(stockist)
            # New group unless same stockist continues
            if groups and _normalize_stockist_key(groups[-1]["stockist_name"]) == key:
                groups[-1]["pages"].append(info)
            else:
                groups.append(
                    {
                        "stockist_name": stockist,
                        "stockist_key": key,
                        "pages": [info],
                    }
                )
        else:
            if groups:
                groups[-1]["pages"].append(info)
            else:
                # Orphan page before first header — keep as unknown segment
                groups.append(
                    {
                        "stockist_name": f"Statement_{info['page_index'] + 1}",
                        "stockist_key": f"STATEMENT_{info['page_index'] + 1}",
                        "pages": [info],
                    }
                )
    return groups


def _extract_statement_from_pdf_group(
    group: Dict[str, Any], filename: str
) -> Dict[str, Any]:
    """Extract one stockist statement from its page group (text and/or images)."""
    pages = group["pages"]
    page_nos = [p["page_index"] + 1 for p in pages]
    combined_text = "\n\n".join(p.get("text") or "" for p in pages).strip()

    result: Optional[Dict[str, Any]] = None

    # Prefer structuring combined OCR/embedded text first (fast, no quota)
    if len(re.sub(r"\s+", "", combined_text)) >= 80:
        result = _structure_sales_text(combined_text, filename, "pdf")
        if result.get("line_items"):
            result["totals"]["extra"]["extraction_method"] = (
                result.get("totals", {}).get("extra", {}).get("extraction_method")
                or "pdf_split_text"
            )

    # Fall back to per-page image extraction and merge
    if not result or not result.get("line_items"):
        merged = empty_result(filename, "pdf")
        merged_items: List[Dict[str, Any]] = []
        for p in pages:
            img = p.get("image_bytes")
            if not img:
                continue
            page_result = _parse_image(
                img, f"{filename}#page{p['page_index'] + 1}", ".png"
            )
            for key in (
                "stockist_name",
                "stockist_address",
                "company_name",
                "period_from",
                "period_to",
                "report_title",
            ):
                if page_result.get(key) and not merged.get(key):
                    merged[key] = page_result[key]
            merged_items.extend(page_result.get("line_items") or [])
            totals = page_result.get("totals") or {}
            if totals.get("sales_value") is not None:
                merged["totals"]["sales_value"] = totals.get("sales_value")
            if totals.get("closing_value") is not None:
                merged["totals"]["closing_value"] = totals.get("closing_value")
        merged["line_items"] = merged_items
        merged["totals"]["extra"]["extraction_method"] = "pdf_split_page_images"
        result = merged

    # Ensure stockist name from split detection wins when vision/OCR confuses manufacturer
    detected = group.get("stockist_name")
    if detected and not re.search(r"^Statement_\d+$", detected):
        current = result.get("stockist_name") or ""
        if (
            not current
            or _MANUFACTURER_HINT.search(current)
            or _normalize_stockist_key(current) != _normalize_stockist_key(detected)
        ):
            # Keep manufacturer if it was misplaced into stockist
            if current and _MANUFACTURER_HINT.search(current) and not result.get(
                "company_name"
            ):
                result["company_name"] = current
            result["stockist_name"] = detected

    result["source_file"] = filename
    result["source_format"] = "pdf"
    result["totals"]["extra"]["split_pages"] = page_nos
    result["totals"]["extra"]["stockist_key"] = group.get("stockist_key")

    # Rebuild combined text with footer OCR boost if TOTAL still missing
    if combined_text and not re.search(r"^\s*TOTAL\b.*\d", combined_text, re.I | re.M):
        boosted_parts = []
        for p in pages:
            boosted_parts.append(
                _merge_footer_total_into_text(p.get("text") or "", p.get("image_bytes"))
            )
        combined_text = "\n\n".join(boosted_parts)

    if combined_text:
        result = _apply_total_row_to_result(result, combined_text)
    return result


def _parse_pdf(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """Split multi-stockist PDFs into statements, then extract each."""
    import os

    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise ValueError("PyMuPDF is required for PDF sales statements") from exc

    max_pages = int(os.getenv("SALES_PDF_MAX_PAGES", "20"))
    zoom = float(os.getenv("SALES_PDF_RENDER_ZOOM", "2.0"))

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        page_infos: List[Dict[str, Any]] = []
        for page_index, page in enumerate(doc):
            if page_index >= max_pages:
                break
            embedded = (page.get_text("text") or "").strip()
            image_bytes: Optional[bytes] = None
            text = embedded
            if len(re.sub(r"\s+", "", embedded)) < 40:
                # 3.5x needed so digits like Issue=12 are not read as 2 on dense scans
                text, image_bytes = _ocr_pdf_page_text(page, zoom=max(zoom, 3.5))
            else:
                # Still render for image fallback if needed later
                pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
                image_bytes = pix.tobytes("png")
                text = _merge_footer_total_into_text(text, image_bytes)

            stockist = _detect_stockist_from_page_text(text)
            page_infos.append(
                {
                    "page_index": page_index,
                    "text": text,
                    "image_bytes": image_bytes,
                    "stockist": stockist,
                }
            )

        groups = _group_pdf_pages_by_stockist(page_infos)
        # Deduplicate adjacent identical keys already handled; merge non-adjacent
        # same stockist only if user wants — keep separate segments in page order.

        if not groups:
            return empty_result(filename, "pdf")

        statements = [
            _extract_statement_from_pdf_group(group, filename) for group in groups
        ]

        # Single stockist → keep flat unified schema (backward compatible)
        if len(statements) == 1:
            single = statements[0]
            single["totals"]["extra"]["statement_count"] = 1
            return single

        return {
            "source_file": filename,
            "source_format": "pdf",
            "multi_statement": True,
            "statement_count": len(statements),
            "statements": statements,
        }
    finally:
        doc.close()


def _structure_sales_text(text: str, filename: str, source_format: str) -> Dict[str, Any]:
    """Turn free-form sales-statement text into unified JSON."""
    import os

    # 0) P.S.PHARMACEUTICALS OPENING/RECEIPT/ISSUE/CLOSING (before Gemini)
    ps = _parse_ps_pharma_statement(text, filename)
    if ps and (ps.get("line_items") or ps.get("totals", {}).get("opening_qty") is not None):
        ps["source_format"] = source_format
        return ps

    # 0b) Mahajan-style OpBal/Receipt/Total/Issue/Closing (before Gemini)
    opbal = _parse_opbal_receipt_issue_statement(text, filename, source_format)
    if opbal and opbal.get("line_items"):
        return opbal

    # 1) Local TXT heuristics (works well for Kaveri-like layouts)
    local = _parse_txt(text.encode("utf-8", errors="ignore"), filename)
    local["source_format"] = source_format
    if local.get("line_items"):
        local["totals"]["extra"]["extraction_method"] = "pdf_text_heuristic"
        return _apply_total_row_to_result(local, text)

    # 2) Gemini text structuring
    try:
        from services.vertex_gemini_client import generate_content_via_vertex

        model = os.getenv("VISION_MODEL", "gemini-2.5-flash-lite").strip()
        payload = {
            "contents": [
                {
                    "role": "user",
                    "parts": [
                        {
                            "text": (
                                _SALES_STATEMENT_TEXT_PROMPT
                                + "\n\nDOCUMENT TEXT:\n"
                                + text[:20000]
                            )
                        }
                    ],
                }
            ],
            "generationConfig": {"temperature": 0.1, "maxOutputTokens": 8192},
        }
        response = generate_content_via_vertex(
            model=model, payload=payload, timeout=120
        )
        parsed = _extract_json_object(_gemini_response_text(response))
        if parsed and parsed.get("line_items"):
            result = empty_result(filename, source_format)
            result = _apply_parsed_sales_json(result, parsed)
            result["totals"]["extra"]["extraction_method"] = "pdf_text_gemini"
            return _apply_total_row_to_result(result, text)
    except Exception as exc:
        logger.warning("PDF text Gemini structuring failed: %s", exc)

    # 3) Generic row heuristic parser
    heuristic = _parse_vikash_ocr_text(text, filename, "." + source_format)
    heuristic["source_format"] = source_format
    return _apply_total_row_to_result(heuristic, text)


# ---------------------------------------------------------------------------
# Image (Gemini Vision)
# ---------------------------------------------------------------------------

_SALES_STATEMENT_VISION_PROMPT = """
You extract pharmaceutical STOCK AND SALES / STOCK REPORT statements (NOT tax invoices).

Return ONLY valid JSON (no markdown) with this exact shape:
{
  "stockist_name": string|null,
  "stockist_address": string|null,
  "company_name": string|null,
  "period_from": "YYYY-MM-DD"|null,
  "period_to": "YYYY-MM-DD"|null,
  "report_title": string|null,
  "line_items": [
    {
      "product_code": string|null,
      "product_name": string,
      "packing": string|null,
      "opening_qty": number,
      "receipts_qty": number,
      "sales_qty": number,
      "sales_value": number,
      "closing_qty": number,
      "closing_value": number,
      "extra": {}
    }
  ],
  "totals": {
    "sales_value": number|null,
    "closing_value": number|null,
    "extra": {}
  }
}

Rules:
- Ignore handwritten notes/circles in margins.
- stockist_name = the agency/seller header (e.g. NEW VIKASH MEDICAL AGENCY), NOT the manufacturer.
- company_name = manufacturer/division line (e.g. VERITAZ HEALTHCARE LTD).
- Dates like 01/06/26 mean DD/MM/YY (year 26 -> 2026). Never invent years like 2001 or 2030.
- Map Pur.Qnt / Purchase / Receipts -> receipts_qty
- Map Sl.Qnt / Sales qty / Issue -> sales_qty; Sl.Value -> sales_value
- Map Cl.Qnt / Closing -> closing_qty; Cl.Value -> closing_value
- Map Op.Qnt / Opening / OpBal -> opening_qty
- For OpBal|Receipt|Total|Issue|Closing: sales_qty=Issue (NOT Total); Dump is not closing_value.
- Use 0 for missing numeric fields.
- Include every printed product row.
""".strip()


_SALES_STATEMENT_TEXT_PROMPT = """
You extract pharmaceutical STOCK AND SALES / STOCK REPORT statements from OCR text (NOT tax invoices).

Return ONLY valid JSON (no markdown) with this exact shape:
{
  "stockist_name": string|null,
  "stockist_address": string|null,
  "company_name": string|null,
  "period_from": "YYYY-MM-DD"|null,
  "period_to": "YYYY-MM-DD"|null,
  "report_title": string|null,
  "line_items": [
    {
      "product_code": string|null,
      "product_name": string,
      "packing": string|null,
      "opening_qty": number,
      "receipts_qty": number,
      "sales_qty": number,
      "sales_value": number,
      "closing_qty": number,
      "closing_value": number,
      "extra": {}
    }
  ],
  "totals": {
    "sales_value": number|null,
    "closing_value": number|null,
    "extra": {}
  }
}

Rules:
- Ignore handwritten notes.
- stockist_name = agency/seller header; company_name = manufacturer/division.
- Dates like 01/06/26 mean DD/MM/YY (year 26 -> 2026).
- Map Op.Qnt/Opening/OpBal -> opening_qty; Pur.Qnt/Purchase/Receipt -> receipts_qty;
  Sl.Qnt/Sales/Issue -> sales_qty; Sl.Value/Amount -> sales_value;
  Cl.Qnt/Closing Balance -> closing_qty; Cl.Value -> closing_value.
- For OpBal | Receipt | Total | Issue | Closing columns:
  opening_qty=OpBal, receipts_qty=Receipt, sales_qty=Issue (NOT Total),
  closing_qty=Closing. Total/Dump/NearExpiry go in extra only; never map Total to sales_qty
  or Dump to closing_value. These are qty-only (sales_value=0, totals money=null).
- Many stockist statements are QTY-ONLY (OpBal/Receipt/Total/Issue/Closing).
  For qty-only reports set sales_value=0 on lines and totals.sales_value=null.
  NEVER invent a huge sales_value from a TOTAL row of quantities.
- TOTAL rows like "TOTAL 107256 81014 188271 75374 123024 35426" are quantities,
  not rupees. Do not put those in totals.sales_value / totals.closing_value.
- For P.S.PHARMACEUTICALS style headers OPENING RECEIPT ISSUE CLOSING, the footer
  "TOTAL 290022 230450 292807 262432" must be captured as:
  totals.opening_qty=290022, totals.receipts_qty=230450,
  totals.sales_qty=292807 (ISSUE), totals.closing_qty=262432.
- Reject OCR-merged numbers (e.g. 75374123024). Prefer null over garbage money.
- Use 0 for missing numeric fields. Include every product row with numbers.
""".strip()


def _image_mime(ext: str) -> str:
    return {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".bmp": "image/bmp",
        ".tif": "image/tiff",
        ".tiff": "image/tiff",
        ".webp": "image/webp",
    }.get(ext, "image/jpeg")


def _ocr_image_to_text(file_bytes: bytes) -> str:
    """Local Tesseract OCR fallback for sales-statement images."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError as exc:
        raise RuntimeError("pytesseract/Pillow required for image OCR fallback") from exc

    import os

    cmd = os.getenv("TESSERACT_CMD", "").strip()
    if cmd:
        pytesseract.pytesseract.tesseract_cmd = cmd
    elif os.name == "nt":
        win_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        if os.path.exists(win_cmd):
            pytesseract.pytesseract.tesseract_cmd = win_cmd

    image = Image.open(io.BytesIO(file_bytes))
    if image.mode not in ("RGB", "L"):
        image = image.convert("RGB")
    # psm 6 (uniform block) reads dense OpBal/Issue qty tables more reliably than default
    return pytesseract.image_to_string(image, config="--psm 6") or ""


def _gemini_response_text(response) -> str:
    data = response.json() if response else {}
    try:
        return (
            data.get("candidates", [{}])[0]
            .get("content", {})
            .get("parts", [{}])[0]
            .get("text", "")
            or ""
        )
    except Exception:
        return ""


def _apply_parsed_sales_json(
    result: Dict[str, Any], parsed: Dict[str, Any]
) -> Dict[str, Any]:
    for key in (
        "stockist_name",
        "stockist_address",
        "company_name",
        "period_from",
        "period_to",
        "report_title",
    ):
        if parsed.get(key):
            val = parsed[key]
            if key in {"period_from", "period_to"}:
                val = _normalize_date(str(val)) or val
            result[key] = val

    # Common vision mix-up: manufacturer vs stockist agency name
    stockist = (result.get("stockist_name") or "").upper()
    company = (result.get("company_name") or "").upper()
    stockist_looks_mfr = bool(
        re.search(r"HEALTHCARE|PHARMA\s*LTD|LABORATOR|LIMITED", stockist)
    ) and not re.search(r"AGENC|STORES|MEDICO|MEDICAL\s+AGENC", stockist)
    company_looks_stockist = bool(
        re.search(r"AGENC|STORES|MEDICO|MEDICAL\s+AGENC", company)
    )
    if stockist_looks_mfr and company_looks_stockist:
        result["stockist_name"], result["company_name"] = (
            result.get("company_name"),
            result.get("stockist_name"),
        )

    title = result.get("report_title") or ""
    m_range = re.search(
        r"(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})\s*[-–to]+\s*(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})",
        title,
        re.I,
    )
    if m_range:
        pf = _normalize_date(m_range.group(1))
        pt = _normalize_date(m_range.group(2))
        if pf and pt:
            result["period_from"], result["period_to"] = pf, pt
    else:
        for pk in ("period_from", "period_to"):
            raw = result.get(pk)
            if not (isinstance(raw, str) and re.match(r"^\d{4}-\d{2}-\d{2}$", raw)):
                continue
            year = int(raw[:4])
            if 2020 <= year <= 2028:
                continue
            yyyy, mm, dd = raw.split("-")
            fixed = _normalize_date(f"{yyyy[2:]}/{mm}/{dd}")
            if fixed:
                result[pk] = fixed
        pf, pt = result.get("period_from"), result.get("period_to")
        if (
            isinstance(pf, str)
            and isinstance(pt, str)
            and re.match(r"^\d{4}-\d{2}-\d{2}$", pf)
            and re.match(r"^\d{4}-\d{2}-\d{2}$", pt)
            and pf[:4] != pt[:4]
        ):
            _y, mm, dd = pt.split("-")
            aligned = f"{pf[:4]}-{mm}-{dd}"
            try:
                datetime.strptime(aligned, "%Y-%m-%d")
                result["period_to"] = aligned
            except ValueError:
                pass

    items = []
    for raw in parsed.get("line_items") or []:
        if not isinstance(raw, dict):
            continue
        item = empty_line_item()
        item["product_code"] = raw.get("product_code")
        item["product_name"] = _clean_name(str(raw.get("product_name") or ""))
        item["packing"] = raw.get("packing")
        item["opening_qty"] = _to_float(raw.get("opening_qty"))
        item["receipts_qty"] = _to_float(raw.get("receipts_qty"))
        item["sales_qty"] = _to_float(raw.get("sales_qty"))
        item["sales_value"] = _to_float(raw.get("sales_value"))
        item["closing_qty"] = _to_float(raw.get("closing_qty"))
        item["closing_value"] = _to_float(raw.get("closing_value"))
        if isinstance(raw.get("extra"), dict):
            item["extra"] = raw["extra"]
        if item["product_name"]:
            items.append(item)
    result["line_items"] = items

    totals = parsed.get("totals") or {}
    if isinstance(totals, dict):
        result["totals"]["sales_value"] = _to_nullable_float(totals.get("sales_value"))
        result["totals"]["closing_value"] = _to_nullable_float(totals.get("closing_value"))
        if isinstance(totals.get("extra"), dict):
            result["totals"]["extra"].update(totals["extra"])
    return result


def _parse_vikash_ocr_text(ocr_text: str, filename: str, ext: str) -> Dict[str, Any]:
    """Heuristic parse for Vikash-style stock report OCR text."""
    # Prefer OpBal/Issue/Closing mapping when that header is present
    opbal = _parse_opbal_receipt_issue_statement(
        ocr_text, filename, (ext or ".pdf").lstrip(".") or "pdf"
    )
    if opbal and opbal.get("line_items"):
        return opbal

    result = empty_result(filename, ext.lstrip("."))
    lines = [ln.strip() for ln in ocr_text.splitlines() if ln.strip()]
    if lines:
        # Prefer agency-looking header over manufacturer
        for ln in lines[:8]:
            if re.search(r"AGENC|STORES|MEDICO|MEDICAL", ln, re.I):
                result["stockist_name"] = _clean_name(ln)
                break
        if not result["stockist_name"]:
            result["stockist_name"] = _clean_name(lines[0])
    for ln in lines[:15]:
        if re.search(r"HEALTHCARE|PHARMA\s*LTD|VERITAZ|AUROBINDO", ln, re.I):
            if not re.search(r"AGENC|STORES|MEDICO", ln, re.I):
                result["company_name"] = _clean_name(ln)
        if re.search(r"Stock\s+Report|Stock\s+and\s+Sale", ln, re.I):
            result["report_title"] = _clean_name(ln)
            m = re.search(
                r"(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})\s*[-–to]+\s*(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})",
                ln,
                re.I,
            )
            if m:
                result["period_from"] = _normalize_date(m.group(1))
                result["period_to"] = _normalize_date(m.group(2))

    # Product rows: NAME ... numbers (op, pur, sl, sl_val, cl, cl_val) — flexible
    row_re = re.compile(
        r"^([A-Z0-9][A-Z0-9 \-./']+?)\s+(-?\d+(?:\.\d+)?(?:\s+-?\d+(?:\.\d+)?){2,})$",
        re.I,
    )
    items: List[Dict[str, Any]] = []
    for ln in lines:
        if re.search(r"Product\s*Name|Op\.Qnt|COMPANY\s*Total|Qnt->", ln, re.I):
            continue
        m = row_re.match(ln)
        if not m:
            # looser: split trailing numbers
            nums = re.findall(r"-?\d+(?:\.\d+)?", ln)
            if len(nums) < 3:
                continue
            name_part = re.split(r"\s+-?\d", ln, maxsplit=1)[0].strip()
            if len(name_part) < 3 or re.search(r"Total|Page|Email|Phone", name_part, re.I):
                continue
            floats = [_to_float(n) for n in nums]
        else:
            name_part = m.group(1).strip()
            floats = [_to_float(n) for n in m.group(2).split()]

        item = empty_line_item()
        item["product_name"] = _clean_name(name_part)
        # Common Vikash: Op Pur Sl SlVal Cl ClVal
        if len(floats) >= 6:
            item["opening_qty"] = floats[0]
            item["receipts_qty"] = floats[1]
            item["sales_qty"] = floats[2]
            item["sales_value"] = floats[3]
            item["closing_qty"] = floats[4]
            item["closing_value"] = floats[5]
        elif len(floats) == 5:
            item["opening_qty"] = floats[0]
            item["sales_qty"] = floats[1]
            item["sales_value"] = floats[2]
            item["closing_qty"] = floats[3]
            item["closing_value"] = floats[4]
        elif len(floats) == 4:
            item["opening_qty"] = floats[0]
            item["sales_qty"] = floats[1]
            item["closing_qty"] = floats[2]
            item["closing_value"] = floats[3]
        else:
            item["extra"]["raw_numbers"] = floats
        items.append(item)

    for ln in lines:
        if re.search(r"COMPANY\s*Total|^\s*TOTAL\b", ln, re.I):
            raw_nums = re.findall(r"-?\d+(?:\.\d+)?", ln)
            # Drop OCR-merged monsters (>8 digits) before mapping to money totals
            nums = [_to_float(n) for n in raw_nums if len(n.split(".")[0].lstrip("-")) <= 8]
            if len(nums) >= 2:
                candidate_sales = nums[-2] if len(nums) >= 3 else None
                candidate_closing = nums[-1]
                # Qty-only TOTAL rows (OpBal/Receipt/Issue/Closing) must not become money
                if _is_implausible_money(candidate_sales) or (
                    candidate_sales is not None and float(candidate_sales) > 1_000_000
                    and float(candidate_closing or 0) > 0
                    and float(candidate_sales) > float(candidate_closing) * 100
                ):
                    result["totals"]["extra"]["total_row_qtys"] = nums
                else:
                    result["totals"]["sales_value"] = candidate_sales
                    result["totals"]["closing_value"] = candidate_closing
            break

    result["line_items"] = items
    result["totals"]["extra"]["extraction_method"] = "tesseract_heuristic"
    return result


def _parse_image(file_bytes: bytes, filename: str, ext: str) -> Dict[str, Any]:
    """Extract sales statement from image via Gemini Vision, with OCR fallback."""
    import os
    import time

    from services.vertex_gemini_client import (
        GeminiProviderError,
        generate_content_via_vertex,
    )

    result = empty_result(filename, ext.lstrip("."))
    mime = _image_mime(ext)
    b64 = base64.b64encode(file_bytes).decode("ascii")
    model = os.getenv("VISION_MODEL", "gemini-2.5-flash-lite").strip()

    vision_payload = {
        "contents": [
            {
                "role": "user",
                "parts": [
                    {"text": _SALES_STATEMENT_VISION_PROMPT},
                    {"inline_data": {"mime_type": mime, "data": b64}},
                ],
            }
        ],
        "generationConfig": {"temperature": 0.1, "maxOutputTokens": 8192},
    }

    # Retry Gemini Vision briefly on 429/503, then fall back to OCR.
    last_err: Optional[Exception] = None
    for attempt in range(3):
        try:
            response = generate_content_via_vertex(
                model=model, payload=vision_payload, timeout=120
            )
            text = _gemini_response_text(response)
            parsed = _extract_json_object(text)
            if parsed and (parsed.get("line_items") or parsed.get("stockist_name")):
                result = _apply_parsed_sales_json(result, parsed)
                result["totals"]["extra"]["extraction_method"] = "gemini_vision"
                return result
            last_err = ValueError(f"non-JSON vision response: {(text or '')[:200]}")
        except GeminiProviderError as exc:
            last_err = exc
            time.sleep(min(2 ** attempt, 8))
        except Exception as exc:
            last_err = exc
            logger.warning("Gemini vision failed for %s: %s", filename, exc)
            break

    logger.warning(
        "Falling back to Tesseract OCR for sales image %s (%s)",
        filename,
        last_err,
    )

    try:
        ocr_text = _ocr_image_to_text(file_bytes)
    except Exception as exc:
        logger.error("Tesseract OCR failed for %s: %s", filename, exc)
        result["totals"]["extra"]["error"] = f"vision_failed: {last_err}; ocr_failed: {exc}"
        return result

    if not ocr_text.strip():
        result["totals"]["extra"]["error"] = "empty OCR text"
        return result

    # Prefer Gemini text structuring of OCR; fall back to heuristic parser.
    text_payload = {
        "contents": [
            {
                "role": "user",
                "parts": [
                    {
                        "text": (
                            _SALES_STATEMENT_TEXT_PROMPT
                            + "\n\nOCR TEXT:\n"
                            + ocr_text[:20000]
                        )
                    }
                ],
            }
        ],
        "generationConfig": {"temperature": 0.1, "maxOutputTokens": 8192},
    }
    try:
        response = generate_content_via_vertex(
            model=model, payload=text_payload, timeout=120
        )
        parsed = _extract_json_object(_gemini_response_text(response))
        if parsed and parsed.get("line_items"):
            result = _apply_parsed_sales_json(result, parsed)
            result["totals"]["extra"]["extraction_method"] = "tesseract_plus_gemini_text"
            return result
    except Exception as exc:
        logger.warning("Gemini text structuring after OCR failed: %s", exc)

    heuristic = _parse_vikash_ocr_text(ocr_text, filename, ext)
    if heuristic.get("line_items"):
        return heuristic

    result["totals"]["extra"]["ocr_preview"] = ocr_text[:2000]
    result["totals"]["extra"]["error"] = "could not structure OCR text"
    return result


def _extract_json_object(text: str) -> Optional[Dict[str, Any]]:
    if not text:
        return None
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)
    try:
        obj = json.loads(cleaned)
        return obj if isinstance(obj, dict) else None
    except json.JSONDecodeError:
        pass
    m = re.search(r"\{.*\}", cleaned, re.DOTALL)
    if not m:
        return None
    try:
        obj = json.loads(m.group(0))
        return obj if isinstance(obj, dict) else None
    except json.JSONDecodeError:
        return None
